import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_left_border(cell, border_hex, size="36"):
    # Add a thick left border to a cell for callout styling
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), border_hex)
    tcBorders.append(left)
    
    # clear other borders
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

# Create document
doc = docx.Document()

# Adjust margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Set base style (Segoe UI for premium styling)
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)
font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B) # Elegant dark gray

# Add elegant header/banner title
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(15)
title_run = title_p.add_run("COMMUNICATION CLIENT\nMigration AssurCore vers Odoo 17")
title_run.bold = True
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Deep Corporate Blue

# Add a divider line
p_div = doc.add_paragraph()
p_div.paragraph_format.space_after = Pt(15)
run_div = p_div.add_run("—" * 60)
run_div.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)

# Add metadata block
p_meta = doc.add_paragraph()
p_meta.paragraph_format.space_after = Pt(15)
p_meta.paragraph_format.line_spacing = 1.2
run_meta = p_meta.add_run("Destinataire : ")
run_meta.bold = True
p_meta.add_run("Zied KAMOUN\n")
run_subj = p_meta.add_run("Objet : ")
run_subj.bold = True
p_meta.add_run("🚀 [Migration Réussie] Bienvenue sur votre nouvel espace AssurCore (Odoo 17) !")

# Add salutation
p_greet = doc.add_paragraph()
p_greet.paragraph_format.space_before = Pt(10)
p_greet.paragraph_format.space_after = Pt(10)
run_greet = p_greet.add_run("Bonjour Zied,")
run_greet.bold = True
run_greet.font.size = Pt(12)

# Add body paragraphs
p_body1 = doc.add_paragraph()
p_body1.paragraph_format.space_after = Pt(10)
p_body1.paragraph_format.line_spacing = 1.15
p_body1.add_run(
    "C'est avec un grand plaisir que nous vous annonçons la finalisation réussie de la migration de votre plateforme "
)
r_bold = p_body1.add_run("AssurCore")
r_bold.bold = True
p_body1.add_run(" depuis l'ancien système Oracle vers la version moderne d'")
r_bold2 = p_body1.add_run("Odoo 17")
r_bold2.bold = True
p_body1.add_run(
    ". Toutes vos données historiques (clients, polices, quittances, règlements, sinistres) ont été préservées avec soin et votre outil dispose désormais de toutes les fonctionnalités avancées d'Odoo, personnalisées spécifiquement pour le marché de l'assurance tunisienne."
)

p_body2 = doc.add_paragraph()
p_body2.paragraph_format.space_after = Pt(15)
p_body2.paragraph_format.line_spacing = 1.15
p_body2.add_run(
    "Vous pouvez dès à présent tester l'application en ligne sur notre serveur sécurisé de production."
)

# Add Callout Box for Credentials
table = doc.add_table(rows=1, cols=1)
table.autofit = False
table.columns[0].width = Inches(6.5)
cell = table.cell(0, 0)
set_cell_background(cell, "F0F4F8") # Soft Light Blue/Gray
set_cell_left_border(cell, "1B365D", size="36") # Thick Deep Blue Left Border

# Add credentials inside the callout box
cell_p = cell.paragraphs[0]
cell_p.paragraph_format.left_indent = Inches(0.15)
cell_p.paragraph_format.right_indent = Inches(0.15)
cell_p.paragraph_format.space_before = Pt(8)
cell_p.paragraph_format.space_after = Pt(4)
r_title = cell_p.add_run("🔐 Vos Informations de Connexion")
r_title.bold = True
r_title.font.size = Pt(12)
r_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

cell_p2 = cell.add_paragraph()
cell_p2.paragraph_format.left_indent = Inches(0.15)
cell_p2.paragraph_format.right_indent = Inches(0.15)
cell_p2.paragraph_format.space_after = Pt(2)
r_label = cell_p2.add_run("•  URL de l'application : ")
r_label.bold = True
cell_p2.add_run("https://assurcore.metadidomi.com")

cell_p3 = cell.add_paragraph()
cell_p3.paragraph_format.left_indent = Inches(0.15)
cell_p3.paragraph_format.right_indent = Inches(0.15)
cell_p3.paragraph_format.space_after = Pt(2)
r_label2 = cell_p3.add_run("•  Identifiant / E-mail : ")
r_label2.bold = True
cell_p3.add_run("borchani0503@gmail.com")

cell_p4 = cell.add_paragraph()
cell_p4.paragraph_format.left_indent = Inches(0.15)
cell_p4.paragraph_format.right_indent = Inches(0.15)
cell_p4.paragraph_format.space_after = Pt(8)
r_label3 = cell_p4.add_run("•  Mot de passe : ")
r_label3.bold = True
r_val3 = cell_p4.add_run("[Votre mot de passe personnalisé]")
r_val3.italic = True

# Add manual section heading
h_man = doc.add_paragraph()
h_man.paragraph_format.space_before = Pt(25)
h_man.paragraph_format.space_after = Pt(10)
r_h = h_man.add_run("📖 Guide de Prise en Main & Test (Mini-Manuel)")
r_h.bold = True
r_h.font.size = Pt(14)
r_h.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

# Add guide paragraphs and items
def add_guide_item(doc, title, description):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r_t = p.add_run(title + " : ")
    r_t.bold = True
    p.add_run(description)

add_guide_item(doc, "1. Découverte de l'Interface", "Connectez-vous et accédez à l'application AssurCore via le menu des applications en haut à gauche. Vous découvrirez un tableau de bord dynamique affichant en temps réel vos indicateurs clés (primes émises, quittances à recouvrer, sinistres en cours).")
add_guide_item(doc, "2. Gestion des Clients & Contacts", "Consultez le menu 'Contacts' ou sous 'AssurCore > Clients'. Tous vos contacts migrés y figurent déjà. Vous pouvez tester la fluidité d'enregistrement d'une nouvelle fiche.")
add_guide_item(doc, "3. Gestion Spécifique de l'Automobile (Risques Tunisiens)", "La gestion complexe de transport interne d'Odoo a été masquée au profit d'un parcours épuré pour le courtage. Lors de la création d'une police 'AUTO', vous pouvez enregistrer ou associer un véhicule avec nos nouveaux champs spécifiques tunisiens : puissance fiscale (CV), valeur vénale (TND), valeur à neuf (TND) et usage du véhicule (Privé, Professionnel, etc.). Le menu d'administration complet des véhicules se trouve sous 'Configuration/Paramétrage > Véhicules'.")
add_guide_item(doc, "4. Suivi des Polices & Quittances", "Naviguez dans les sections 'Polices' et 'Quittances' pour suivre les contrats actifs, les états de recouvrement (émise, payée, soldée) et tester l'enregistrement simplifié des règlements.")
add_guide_item(doc, "5. Rapports & Extraction Excel", "Testez la génération de rapports dynamiques sous 'Rapports > Mouvement Clients'. Vous pourrez y filtrer vos données historiques et exporter un rapport propre et structuré sous format Microsoft Excel (.xlsx).")

# Add outro
p_outro = doc.add_paragraph()
p_outro.paragraph_format.space_before = Pt(15)
p_outro.paragraph_format.space_after = Pt(10)
p_outro.paragraph_format.line_spacing = 1.15
p_outro.add_run(
    "Nous sommes convaincus que cette nouvelle interface moderne et ces performances accrues transformeront positivement votre gestion quotidienne. Nous restons à votre entière disposition pour recueillir vos retours de tests."
)

p_sign = doc.add_paragraph()
p_sign.paragraph_format.space_before = Pt(20)
p_sign.paragraph_format.space_after = Pt(0)
r_sign1 = p_sign.add_run("Bien cordialement,\n\n")
r_sign2 = p_sign.add_run("L'Équipe Projet AssurCore")
r_sign2.bold = True

doc.save("d:\\Robot\\ASSURPROD\\Mail_Migration_AssurCore.docx")
print("Word document generated successfully at d:\\Robot\\ASSURPROD\\Mail_Migration_AssurCore.docx")
